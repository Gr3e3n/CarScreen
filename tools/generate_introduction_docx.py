import os
from datetime import datetime


def main() -> None:
    """
    Generate a comprehensive introduction.docx into workspace root.

    Dependencies:
      - python-docx
      - matplotlib (for diagram PNGs embedded in docx)
    """
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            "缺少依赖 python-docx。请先执行：python -m pip install python-docx\n"
            f"原始错误：{e}"
        )

    try:
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            "缺少依赖 matplotlib。请先执行：python -m pip install matplotlib\n"
            f"原始错误：{e}"
        )

    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    diagrams_dir = os.path.join(project_root, "tools", "_diagrams")
    os.makedirs(diagrams_dir, exist_ok=True)

    # ── Diagram helpers ───────────────────────────────────────────────

    def _setup_ax(ax, title: str) -> None:
        ax.set_title(title, fontsize=14, fontweight="bold")
        ax.set_xlim(0, 100)
        ax.set_ylim(0, 100)
        ax.axis("off")

    def _box(ax, x, y, w, h, text, fc="#F8FAFC", ec="#CBD5E1", lw=1.5) -> None:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.6,rounding_size=6",
            linewidth=lw,
            edgecolor=ec,
            facecolor=fc,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=10)

    def _arrow(ax, x1, y1, x2, y2, text=None, color="#64748B") -> None:
        arr = FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.4,
            color=color,
        )
        ax.add_patch(arr)
        if text:
            ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 2, text, ha="center", va="center", fontsize=9, color=color)

    def gen_arch_overview_png(path: str) -> None:
        fig, ax = plt.subplots(figsize=(12.5, 7.2), dpi=200)
        _setup_ax(ax, "CarScreen Architecture Overview (Read-only Display/Judgement)")

        # Columns
        _box(ax, 3, 92, 30, 6, "UI Layer", fc="#E0F2FE", ec="#0284C7")
        _box(ax, 35, 92, 30, 6, "Domain/Logic", fc="#ECFDF5", ec="#10B981")
        _box(ax, 67, 92, 30, 6, "Infrastructure", fc="#FFF7ED", ec="#F97316")

        # UI
        _box(ax, 3, 78, 30, 10, "MainActivity\n(Safety Center)")
        _box(ax, 3, 66, 30, 10, "ModuleDetailActivity\n(Read-only detail)")
        _box(ax, 3, 52, 30, 10, "AccidentTraceActivity\n(Event list)")
        _box(ax, 3, 40, 30, 10, "AccidentTraceDetailActivity\n(EDR + Evidence)")
        _box(ax, 3, 27, 30, 10, "Adapters + XML/ViewBinding\n(ModuleAdapter, layouts)")

        # Domain
        _box(ax, 35, 78, 30, 10, "SafetyViewModel\n(LiveData modules)")
        _box(ax, 35, 66, 30, 10, "MockDataProvider\n+ risk rules")
        _box(ax, 35, 54, 30, 10, "ModuleCatalog\n(param mapping)")
        _box(ax, 35, 40, 30, 10, "AccidentTraceViewModel")
        _box(ax, 35, 28, 30, 10, "AccidentRepository\n(LiveData + evidence)")
        _box(ax, 35, 16, 30, 10, "ResponsibilityAnalyzer\n+ Sha256")

        # Infra
        _box(ax, 67, 78, 30, 10, "AccidentMonitorService\n(Foreground service)")
        _box(ax, 67, 66, 30, 10, "AccidentMonitor\n(20Hz + RingBuffer)")
        _box(ax, 67, 52, 30, 10, "Room DB\n(Database/Dao/Entities)")
        _box(ax, 67, 40, 30, 10, "Resources\n(drawable/colors)")

        # Arrows
        _arrow(ax, 18, 78, 35, 83, "observe")
        _arrow(ax, 18, 66, 35, 71, "read-only")
        _arrow(ax, 18, 52, 35, 45, "observe")
        _arrow(ax, 18, 40, 35, 33, "load detail")
        _arrow(ax, 50, 78, 50, 74, None)
        _arrow(ax, 50, 66, 50, 62, None)
        _arrow(ax, 50, 54, 50, 50, None)
        _arrow(ax, 82, 78, 82, 74, None)
        _arrow(ax, 82, 66, 82, 62, None)
        _arrow(ax, 82, 66, 65, 33, "upsert")
        _arrow(ax, 82, 66, 82, 56, "persist")
        _arrow(ax, 82, 52, 65, 33, "query")
        _arrow(ax, 65, 28, 65, 22, "analyze/hash")

        ax.text(50, 6, "Note: current version does NOT send control commands; only displays/infers risk.",
                ha="center", va="center", fontsize=10, color="#334155")

        fig.tight_layout()
        fig.savefig(path, bbox_inches="tight")
        plt.close(fig)

    def gen_safety_flow_png(path: str) -> None:
        fig, ax = plt.subplots(figsize=(12.5, 4.8), dpi=200)
        _setup_ax(ax, "Safety Center Flow (Navigation + Data)")
        _box(ax, 5, 60, 22, 18, "MainActivity\nRecyclerView")
        _box(ax, 32, 60, 22, 18, "SafetyViewModel\nLiveData")
        _box(ax, 59, 60, 22, 18, "MockDataProvider\n(status+risk)")
        _box(ax, 32, 25, 22, 18, "ModuleCatalog\n(param list)")
        _box(ax, 59, 25, 22, 18, "ModuleDetailActivity\n(read-only)")
        _arrow(ax, 27, 69, 32, 69, "observe")
        _arrow(ax, 43, 60, 43, 43, "moduleId")
        _arrow(ax, 54, 34, 59, 34, "render params")
        _arrow(ax, 54, 69, 59, 69, "refresh list")
        _arrow(ax, 16, 60, 70, 34, "tap/long-press")
        fig.tight_layout()
        fig.savefig(path, bbox_inches="tight")
        plt.close(fig)

    def gen_edr_sequence_png(path: str) -> None:
        fig, ax = plt.subplots(figsize=(12.5, 5.2), dpi=200)
        _setup_ax(ax, "EDR Sequence (Foreground Service + 20Hz Sampling + Persist)")
        _box(ax, 5, 80, 18, 10, "MainActivity")
        _box(ax, 30, 80, 22, 10, "AccidentMonitorService")
        _box(ax, 58, 80, 18, 10, "AccidentMonitor")
        _box(ax, 80, 80, 15, 10, "Room DB")

        # timeline y levels
        ax.plot([10, 10], [15, 78], color="#94A3B8", linewidth=1)
        ax.plot([41, 41], [15, 78], color="#94A3B8", linewidth=1)
        ax.plot([67, 67], [15, 78], color="#94A3B8", linewidth=1)
        ax.plot([87, 87], [15, 78], color="#94A3B8", linewidth=1)

        # messages
        _arrow(ax, 10, 70, 41, 70, "startForegroundService()")
        _arrow(ax, 41, 62, 67, 62, "start(dao)")
        _arrow(ax, 67, 54, 67, 54, "20Hz loop + RingBuffer")
        _arrow(ax, 67, 46, 67, 46, "trigger -> freeze windows")
        _arrow(ax, 67, 38, 87, 38, "persist events/telemetry/resp")
        _arrow(ax, 67, 30, 41, 30, "upsertCapturedEvent()")
        _arrow(ax, 41, 22, 10, 22, "UI observes LiveData")

        ax.text(50, 10, "Trigger conditions: strong decel / emergency brake / rare demo trigger",
                ha="center", va="center", fontsize=10, color="#334155")
        fig.tight_layout()
        fig.savefig(path, bbox_inches="tight")
        plt.close(fig)

    def gen_risk_rules_png(path: str) -> None:
        fig, ax = plt.subplots(figsize=(12.5, 4.8), dpi=200)
        _setup_ax(ax, "Risk Judgement Flow (Read-only)")
        _box(ax, 6, 60, 24, 16, "Input\n(module.status)")
        _box(ax, 38, 60, 24, 16, "Rule Engine\n(MockDataProvider)")
        _box(ax, 70, 60, 24, 16, "Output\nriskLevel + reason")
        _box(ax, 38, 28, 24, 16, "UI Rendering\n(ModuleAdapter/Detail)")
        _arrow(ax, 30, 68, 38, 68, "parse/keyword/threshold")
        _arrow(ax, 62, 68, 70, 68, "classify")
        _arrow(ax, 82, 60, 50, 44, "display")
        _arrow(ax, 30, 60, 50, 44, "tap -> detail")
        ax.text(50, 18, "No control commands. Only display/judgement based on readable parameters.",
                ha="center", va="center", fontsize=10, color="#334155")
        fig.tight_layout()
        fig.savefig(path, bbox_inches="tight")
        plt.close(fig)

    def generate_diagrams() -> dict[str, str]:
        files = {
            "arch": os.path.join(diagrams_dir, "01_arch_overview.png"),
            "safety_flow": os.path.join(diagrams_dir, "02_safety_flow.png"),
            "edr_seq": os.path.join(diagrams_dir, "03_edr_sequence.png"),
            "risk_flow": os.path.join(diagrams_dir, "04_risk_flow.png"),
        }
        gen_arch_overview_png(files["arch"])
        gen_safety_flow_png(files["safety_flow"])
        gen_edr_sequence_png(files["edr_seq"])
        gen_risk_rules_png(files["risk_flow"])
        return files

    diagrams = generate_diagrams()

    # ── Doc helpers ───────────────────────────────────────────────────

    doc = Document()

    def h1(text: str) -> None:
        doc.add_heading(text, level=1)

    def h2(text: str) -> None:
        doc.add_heading(text, level=2)

    def h3(text: str) -> None:
        doc.add_heading(text, level=3)

    def p(text: str) -> None:
        doc.add_paragraph(text)

    def bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def add_picture(path: str, caption: str) -> None:
        doc.add_picture(path, width=Inches(6.5))
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True if cap.runs else False

    def add_table(headers: list[str], rows: list[list[str]]) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
        for r in rows:
            cells = table.add_row().cells
            for i, v in enumerate(r):
                cells[i].text = v

    # ── Content ───────────────────────────────────────────────────────

    h1("CarScreen 工程全解（增强版：含图/表/时序/类关系）")
    p(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    p(f"工程目录：{project_root}")

    h2("0. 阅读指南（你能从本文得到什么）")
    bullet("快速理解：本工程展示哪些界面/功能、对应哪些参数、如何形成完整的“事故溯源+可信存证+责任界定”链路。")
    bullet("落地定位：每个关键代码文件在哪里、负责什么、由谁调用、数据怎么流转。")
    bullet("可扩展点：未来接入真实车辆参数订阅（只读/可读写均可）的最小改动路径。")

    h2("1. 工程目标与约束（最重要的设计前提）")
    bullet("目标：做一个车机端“智能安全中心”与“事故溯源（EDR）”演示工程，强调“可解释 + 可追溯 + 可存证”。")
    bullet("约束：当前版本严格遵守“只能读取参数用于展示/判断”，UI 不下发控制命令。")
    bullet("现状：安全中心状态来自 Mock（每 2 秒刷新）；事故溯源采样/触发/写库为模拟数据，但链路结构已按真实 EDR 设计。")

    h2("2. 架构总览（图）")
    p("下图是本工程的层次划分与关键依赖关系（UI/业务/基础设施）。")
    add_picture(diagrams["arch"], "图 1：架构总览（只读展示/判断，不下发控制）")

    h2("3. 功能清单（工程对外展示了什么）")
    h3("3.1 智能安全中心")
    bullet("首页模块卡片：展示各安全能力状态，并在必要时打“关注/高风险”标签，原因可见。")
    bullet("模块详情页：展示该模块的只读参数映射（propertyId/valueType/meaning）、当前概览、风险结论与建议。")
    h3("3.2 事故溯源（EDR）")
    bullet("溯源列表：事件列表（Room 历史 + 运行时新触发合并，自动刷新）。")
    bullet("溯源详情：事故基本信息、事故前 10 秒关键遥测、可解释责任指标、责任占比、系统链路/环境快照、可信存证摘要。")

    h2("4. 关键流程（图 + 解释）")
    h3("4.1 安全中心导航与数据流")
    add_picture(diagrams["safety_flow"], "图 2：安全中心流程（导航 + 数据）")
    bullet("MainActivity 仅负责展示与导航，不负责规则计算。")
    bullet("SafetyViewModel 提供 modules LiveData，当前由 MockDataProvider 生成；未来可替换为真实订阅更新。")
    bullet("ModuleCatalog 是参数映射“单一事实源”（哪个模块展示哪些 propertyId、含义是什么）。")

    h3("4.2 事故溯源：前台服务 + 20Hz 采样 + 冻结窗口 + Room 持久化")
    add_picture(diagrams["edr_seq"], "图 3：事故溯源时序（前台服务保活 + 20Hz 采样 + 写库 + UI 刷新）")
    bullet("核心思想：高频采样不直接写库，而是先写环形缓冲；触发后再冻结窗口并写入 Room。")
    bullet("这样既能保留事故前关键数据，又能避免长期高频 IO。")

    h3("4.3 风险判定流程（只读）")
    add_picture(diagrams["risk_flow"], "图 4：风险判定流程（基于只读状态做分类与解释）")
    bullet("风险规则放在 MockDataProvider（演示版）；真实接入建议迁移到独立 RuleEngine，便于测试与复用。")

    h2("5. 文件职责表（表格）")
    add_table(
        headers=["文件路径", "职责", "关键被谁调用/依赖"],
        rows=[
            ["app/src/main/java/com/example/carscreen/MainActivity.kt", "安全中心首页：启动服务、展示模块、导航到详情/溯源", "启动 AccidentMonitorService；观察 SafetyViewModel.modules"],
            ["app/src/main/java/com/example/carscreen/ModuleDetailActivity.kt", "模块详情页：参数只读展示、风险结论展示", "读取 ModuleCatalog；观察 SafetyViewModel.modules"],
            ["app/src/main/java/com/example/carscreen/SafetyViewModel.kt", "模块数据提供者：当前为 Mock 刷新", "调用 MockDataProvider.generateModules()"],
            ["app/src/main/java/com/example/carscreen/mock/MockDataProvider.kt", "Mock 状态生成 + 风险规则（riskLevel/riskReason）", "供 SafetyViewModel 使用"],
            ["app/src/main/java/com/example/carscreen/data/ModuleCatalog.kt", "模块参数映射目录（propertyId/valueType/meaning/建议）", "供 ModuleDetailActivity 使用"],
            ["app/src/main/java/com/example/carscreen/service/AccidentMonitorService.kt", "前台服务：保活监控与采样循环", "在 MainActivity 启动；内部启动 AccidentMonitor"],
            ["app/src/main/java/com/example/carscreen/trace/AccidentMonitor.kt", "20Hz 采样、环形缓冲、触发冻结、写 Room、更新仓库", "由 AccidentMonitorService 启动"],
            ["app/src/main/java/com/example/carscreen/trace/AccidentRepository.kt", "内存 LiveData + Room 合并；生成环境/链路；生成存证", "被 AccidentMonitor 与 ViewModel/DetailActivity 调用"],
            ["app/src/main/java/com/example/carscreen/trace/ResponsibilityAnalyzer.kt", "可解释指标计算 + 责任推断（可解释）", "被 AccidentMonitor/DetailActivity 调用"],
            ["app/src/main/java/com/example/carscreen/db/*", "Room 数据层：entities/dao/db", "被 AccidentMonitorService/Repository 使用"],
            ["app/src/main/res/layout/*", "页面与组件布局（XML + ViewBinding）", "被各 Activity/Adapter 使用"],
        ],
    )

    h2("6. 模块参数映射表（表格，来自 info.txt）")
    add_table(
        headers=["模块ID", "模块名", "参数（NAME=property）", "propertyId", "valueType", "用途"],
        rows=[
            ["adas", "驾驶辅助", "AUTONOMOUS_EMERGENCY_BRAKING", "24320", "boolean", "AEB 开关状态展示/风险提示"],
            ["adas", "驾驶辅助", "LANE_DEPARTURE_WARNING", "43776", "boolean", "偏离预警状态展示"],
            ["adas", "驾驶辅助", "LANE_KEEPING_AID", "60928", "boolean", "车道保持状态展示"],
            ["collision", "碰撞预警", "FORWARD_COLLISION_WARN_SNVT", "29184", "int", "灵敏度/是否关闭的风险提示"],
            ["blindspot", "盲区监测", "LANE_CHANGE_WARNING_MODE", "61696", "int", "提示方式展示（视觉/声音等）"],
            ["fatigue", "疲劳监测", "DMS_DRIVER_FATIGUE_STATUS", "93952", "int", "疲劳/分心风险提示"],
            ["door", "车门状态", "DOOR_OPEN_WARN_ACTIVE", "29696", "boolean", "车门提醒开关展示"],
            ["speed_limit", "限速提醒", "SPEED_LIMIT_WARNING_MODE", "115456", "int", "限速提醒模式展示"],
            ["speed_limit", "限速提醒", "SPEED_LIMIT_WARNING_OFFSET_VALUE", "122112", "int", "偏差值大小提示（越大风险越高）"],
            ["rain_safety", "雨天安全", "AUTO_CLOSE_WINDOW_RAINY", "23296", "boolean", "雨天自动关窗状态展示"],
            ["rain_safety", "雨天安全", "AUTO_REAR_WIPING", "122880", "boolean", "倒车自动后雨刮状态展示"],
            ["rain_safety", "雨天安全", "LOCK_AUTO_CLOSE_WINDOW", "23040", "boolean", "锁车自动关窗状态展示"],
            ["child_safety", "乘员安全", "CHILD_SAFETY_LOCK", "35328", "boolean", "儿童锁状态（info 标记 deprecated）"],
            ["child_safety", "乘员安全", "PAB_SWITCH", "143872", "boolean", "副驾安全气囊使能状态展示"],
        ],
    )

    h2("7. 风险规则表（表格）")
    add_table(
        headers=["模块ID", "规则触发条件（基于 status 文本/阈值）", "风险等级", "reason 示例"],
        rows=[
            ["fatigue", "status 含“疲劳”", "高风险", "检测到疲劳驾驶"],
            ["fatigue", "status 含“分心/未知”", "关注", "检测到分心驾驶 / 状态未知"],
            ["door", "status 含“开启”且非“全部关闭”", "关注", "存在车门未关"],
            ["collision", "status 含“关”", "关注", "碰撞预警可能处于关闭状态"],
            ["lane", "status 含“关”", "关注", "车道偏离预警关闭"],
            ["lane_keep", "status 含“关闭”", "关注", "车道保持辅助关闭"],
            ["speed_limit", "abs(offset)>=15", "高风险", "限速偏差过大"],
            ["speed_limit", "abs(offset)>=10", "关注", "限速偏差偏大"],
            ["rain_safety", "雨天关窗关/锁车关窗关 任一", "关注", "雨天自动关窗关闭等"],
            ["child_safety", "儿童锁关", "关注", "儿童锁未开启"],
        ],
    )

    h2("8. 事故溯源触发与持久化（表格）")
    add_table(
        headers=["阶段", "关键类/函数", "核心逻辑", "输出/落库"],
        rows=[
            ["保活启动", "AccidentMonitorService.onCreate", "创建通知渠道并 startForeground；启动 AccidentMonitor", "服务常驻运行"],
            ["高频采样", "AccidentMonitor.runSamplingLoop", "每 50ms sampleTelemetry 写 RingBuffer", "内存环形缓冲"],
            ["触发判断", "AccidentMonitor.shouldTrigger", "强减速/紧急制动/低概率演示触发", "进入冻结窗口流程"],
            ["冻结窗口", "AccidentMonitor.triggerCapture", "冻结事故前 10s；采集事故后 10s", "AccidentDetailBundle"],
            ["责任分析", "ResponsibilityAnalyzer.analyze/inferResponsibility", "计算反应时间/制动上升/TTC/AEB等 + 推断占比", "ResponsibilityResult + metrics"],
            ["持久化", "AccidentMonitor.persistToRoom", "写入 event/telemetry/responsibility", "Room：accident_edr.db"],
            ["UI刷新", "AccidentRepository.upsertCapturedEvent", "更新 LiveData 事件列表", "溯源列表实时更新"],
            ["存证", "AccidentRepository.generateEvidence", "payload -> sha256 -> 模拟签名/txid -> 入库", "Room：evidence_records"],
        ],
    )

    h2("9. 重要实现细节（我认为关键且容易踩坑的点）")
    h3("9.1 为什么用“环形缓冲 + 冻结窗口”而不是持续写库？")
    p("EDR 的价值在于“事故前的关键窗口”。持续 20Hz 写库会导致 IO 压力与存储增长；环形缓冲能常态轻量运行，触发后再把关键窗口一次性落库。")
    h3("9.2 为什么要前台服务？")
    p("车机/Android 系统可能在后台回收进程；前台服务是保证采样链路持续运行的最小可用手段。")
    h3("9.3 责任分析为什么强调“可解释”？")
    p("保险/争议场景里，黑盒结论很难被接受。本工程将反应时间、制动上升、峰值减速、TTC、AEB介入等指标显式计算并展示，便于复核。")
    h3("9.4 “只读展示/判断”如何保证不会误触发控制？")
    p("交互层统一改为跳转详情页；保留的 sendControlCommand 未在 UI 上触发（可后续按权限与车型能力再开放）。")

    h2("10. 后续接入真实参数（建议路径）")
    bullet("安全中心：将 SafetyViewModel 的 Mock 轮询替换为真实订阅更新，保持 SafetyModule 与风险渲染不变。")
    bullet("事故溯源：将 AccidentMonitor.sampleTelemetry() 替换为真实车辆总线/属性订阅值；触发阈值按车型标定。")
    bullet("存证：用真实链服务 SDK 替换模拟 txid/signature，并将 payload 规范化（字段顺序/单位/签名算法）。")

    h2("11. 简单陈述（OK 段落）")
    p("OK：当前工程已形成可运行的端到端演示闭环——安全中心负责“状态展示与风险提示”，事故溯源负责“采样→触发→冻结→分析→存证→追溯”。在“只读”约束下，系统仍能提供可解释的安全态势与事故证据链框架；后续接入真实车辆参数后，可在不改变 UI 结构的前提下逐步增强准确性与覆盖面。")

    out_path = os.path.join(project_root, "introduction.docx")
    tmp_path = os.path.join(project_root, "introduction.tmp.docx")
    doc.save(tmp_path)
    try:
        if os.path.exists(out_path):
            os.replace(out_path, os.path.join(project_root, "introduction.prev.docx"))
    except PermissionError:
        # If the doc is open in Word, Windows will lock it. Keep tmp and exit gracefully.
        print("注意：introduction.docx 可能正在被 Word 打开导致写入失败。已生成临时文件：")
        print(tmp_path)
        return
    os.replace(tmp_path, out_path)
    # cleanup previous backup if any
    prev = os.path.join(project_root, "introduction.prev.docx")
    if os.path.exists(prev):
        try:
            os.remove(prev)
        except Exception:
            pass
    print(f"已生成：{out_path}")


if __name__ == "__main__":
    main()

